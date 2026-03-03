#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 轉 DOCX/PDF 合併轉換工具
將所有 Markdown 檔案合併並轉換為 DOCX 和 PDF 檔案
"""

import os
import re
import glob
import sys
import traceback
import pypandoc
from docx import Document
from docx.shared import Pt
from docx2pdf import convert as docx2pdf_convert

# Word WdTCSCConverterDirection 常數
# SCTC = 0 = Simplified to Traditional Chinese（簡體→繁體）
# TCSC = 1 = Traditional to Simplified Chinese（繁體→簡體）
# Auto = 2 = 自動判斷
WD_TCSC_CONVERTER_DIRECTION_SCTC = 0

def get_article_number(filename):
    """從檔案名提取文章編號"""
    match = re.search(r'_(\d+)\.md$', filename)
    if match:
        return int(match.group(1))
    return 0

def extract_book_name(filename):
    """從檔案名提取書名（只取第一個底線前的部分）
    範例：
        - 属灵人_卷一 灵魂体的总论_1.md → "属灵人"
        - 神经纶的福音_1.md → "神经纶的福音"
    """
    # 移除 .md 後綴
    name = filename.replace('.md', '')
    
    # 只提取第一個底線前的內容作為書名
    if '_' in name:
        return name.split('_')[0]
    
    return name

def extract_book_and_volume_name(filename):
    """從檔案名提取書名和卷名的組合（用於分組）
    範例：
        - 属灵人_卷一 灵魂体的总论_1.md → "属灵人_卷一 灵魂体的总论"
        - 神经纶的福音_1.md → "神经纶的福音"
    """
    # 移除 .md 後綴
    name = filename.replace('.md', '')
    
    # 分割檔案名
    parts = name.split('_')
    
    if len(parts) >= 3:
        # 三級結構：返回「書名_卷名」
        return f"{parts[0]}_{parts[1]}"
    elif len(parts) == 2:
        # 兩級結構：返回書名
        return parts[0]
    else:
        return name

def chinese_to_number(chinese_str):
    """將中文數字轉換為阿拉伯數字"""
    mapping = {
        '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
        '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
        '百': 100
    }
    
    # 處理簡單情況
    if chinese_str in mapping:
        return mapping[chinese_str]
    
    # 處理十位數（如「十一」、「二十」等）
    result = 0
    temp = 0
    for char in chinese_str:
        if char in mapping:
            if mapping[char] >= 10:
                if temp == 0:
                    temp = 1
                result += temp * mapping[char]
                temp = 0
            else:
                temp = mapping[char]
        else:
            if temp > 0:
                result += temp
                temp = 0
    
    if temp > 0:
        result += temp
    
    return result if result > 0 else 0

def extract_sort_key(filename):
    """提取排序鍵，確保正確排序
    範例：
        - 属灵人_卷一 灵魂体的总论_1.md → (1, 1)
        - 属灵人_卷一 灵魂体的总论_2.md → (1, 2)
        - 属灵人_卷二 肉体_1.md → (2, 1)
        - 神经纶的福音_1.md → (0, 1)
    """
    name = filename.replace('.md', '')
    parts = name.split('_')
    
    if len(parts) >= 3:
        # 三級結構：提取卷號和章號
        volume_text = parts[1]
        chapter_num = int(parts[2]) if parts[2].isdigit() else 0
        
        # 嘗試從卷名中提取數字
        volume_match = re.search(r'卷([一二三四五六七八九十百]+)|第(\d+)[册冊]|第([一二三四五六七八九十]+)[册冊]', volume_text)
        
        if volume_match:
            if volume_match.group(1):  # 中文數字
                volume_num = chinese_to_number(volume_match.group(1))
            elif volume_match.group(2):  # 阿拉伯數字
                volume_num = int(volume_match.group(2))
            elif volume_match.group(3):  # 中文數字（冊）
                volume_num = chinese_to_number(volume_match.group(3))
            else:
                volume_num = 0
            return (volume_num, chapter_num)
        else:
            return (0, chapter_num)
    
    elif len(parts) == 2:
        # 兩級結構：只有章號
        chapter_num = int(parts[1]) if parts[1].isdigit() else 0
        return (0, chapter_num)
    
    return (0, 0)

def collect_markdown_files_by_book():
    """收集所有 Markdown 檔案並按書名和卷名分組（一卷一個 docx）"""
    # 找出所有 *_數字.md 格式的檔案
    all_files = glob.glob("*_*.md")
    
    # 過濾出符合格式的檔案（書名_數字.md 或 書名_卷名_數字.md）
    valid_files = [f for f in all_files if re.match(r'^.+_\d+\.md$', f)]
    
    if not valid_files:
        return {}
    
    # 按「書名_卷名」分組並排序
    books = {}
    for file in valid_files:
        book_and_volume = extract_book_and_volume_name(file)
        if book_and_volume:
            if book_and_volume not in books:
                books[book_and_volume] = []
            books[book_and_volume].append(file)
    
    # 每個分組的檔案按排序鍵排序（支援三級結構）
    for book_and_volume in books:
        books[book_and_volume].sort(key=extract_sort_key)
    
    return books

def process_markdown_headings(content, article_num):
    """處理 Markdown 標題：
    將連續的一級標題合併為一個標題
    保留其他級別標題的原有格式（保持視覺差異）
    """
    lines = content.split('\n')
    processed_lines = []
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 處理一級標題
        if line.strip().startswith('# '):
            # 收集所有連續的一級標題（包括中間有空行的情況）
            h1_titles = []
            j = i
            
            # 收集連續的一級標題
            while j < len(lines):
                current_line = lines[j].strip()
                if current_line.startswith('# '):
                    title_text = current_line[2:].strip()
                    h1_titles.append(title_text)
                    j += 1
                elif current_line == '':
                    # 空行，繼續檢查下一行是否是一級標題
                    j += 1
                    continue
                else:
                    # 遇到非空行且非一級標題，停止收集
                    break
            
            # 處理收集到的一級標題：簡單合併所有標題
            if h1_titles:
                merged_title = ' '.join(h1_titles)
                processed_lines.append(f"# {merged_title}")
            else:
                # 不應該發生，但安全處理
                processed_lines.append(line)
            
            # 跳過已處理的行（包括空行）
            i = j
        else:
            # 保留其他所有內容（包括其他級別標題）不變
            processed_lines.append(line)
            i += 1
    
    return '\n'.join(processed_lines)

def merge_markdown_files(files):
    """合併所有 Markdown 檔案內容"""
    merged_content = []
    
    for i, file in enumerate(files):
        print(f"讀取檔案 [{i+1}/{len(files)}]: {file}")
        
        try:
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read().strip()
                
                # 將波浪號轉義
                content = content.replace('~', r'\~')
                
                # 獲取文章編號
                article_num = get_article_number(file)
                
                # 處理標題
                content = process_markdown_headings(content, article_num)
                
                merged_content.append(content)
                
                # 在每篇文章之間添加分頁符號（除了最後一篇）
                if i < len(files) - 1:
                    merged_content.append('\n\n```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n\n')
        
        except Exception as e:
            print(f"警告：無法讀取檔案 {file}: {e}")
            continue
    
    return '\n\n'.join(merged_content)

def set_line_spacing(docx_file, line_spacing=1.5):
    """設定 DOCX 檔案的行距"""
    try:
        print(f"設定行距為 {line_spacing}...")
        doc = Document(docx_file)
        
        # 遍歷所有段落並設定行距
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = line_spacing
        
        # 遍歷所有表格中的段落並設定行距
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.line_spacing = line_spacing
        
        # 儲存修改
        doc.save(docx_file)
        print(f"行距設定完成")
        return True
        
    except Exception as e:
        print(f"\n警告：設定行距時發生錯誤: {e}")
        return False

def process_docx_headings(docx_file):
    """處理 DOCX 檔案中的標題樣式：
    只保留一級標題（Heading 1）作為標題樣式，其他級別的標題改為自訂格式
    保持視覺差異（字體大小、粗體等），但不會成為 PDF 書籤
    """
    try:
        print("處理標題樣式，確保只有一級標題用於 PDF 書籤...")
        doc = Document(docx_file)
        
        def get_heading_format(style_name):
            """從標題樣式中獲取格式資訊（字體大小、是否粗體）"""
            try:
                style = doc.styles[style_name]
                # 獲取字體大小（如果有的話）
                font_size = None
                is_bold = False
                
                if hasattr(style, 'font') and style.font.size:
                    font_size = style.font.size
                if hasattr(style, 'font') and style.font.bold is not None:
                    is_bold = style.font.bold
                
                # 如果沒有從樣式獲取到，使用預設值
                if font_size is None:
                    # 根據標題級別設定預設字體大小
                    if '2' in style_name or '二' in style_name:
                        font_size = Pt(18)
                    elif '3' in style_name or '三' in style_name:
                        font_size = Pt(16)
                    elif '4' in style_name or '四' in style_name:
                        font_size = Pt(14)
                    elif '5' in style_name or '五' in style_name:
                        font_size = Pt(12)
                    else:
                        font_size = Pt(14)  # 預設值
                
                # 預設標題為粗體
                if is_bold is False and style_name.startswith(('Heading', '標題')):
                    is_bold = True
                
                return font_size, is_bold
            except:
                # 如果無法獲取樣式，使用預設值
                return Pt(14), True
        
        def apply_custom_format(paragraph, original_style_name):
            """將段落改為 Normal 樣式，但保留原有的視覺格式"""
            # 先獲取原樣式的格式資訊
            font_size, is_bold = get_heading_format(original_style_name)
            
            # 改為 Normal 樣式（不會成為 PDF 書籤）
            paragraph.style = doc.styles['Normal']
            
            # 應用原樣式的格式到所有 runs
            for run in paragraph.runs:
                if font_size:
                    run.font.size = font_size
                run.font.bold = is_bold
        
        # 需要處理的非一級標題樣式
        heading_styles_to_process = [
            'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5',
            '標題 2', '標題 3', '標題 4', '標題 5'
        ]
        
        # 遍歷所有段落
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else None
            
            # 如果段落使用的是非一級標題樣式，改為自訂格式
            if style_name in heading_styles_to_process:
                apply_custom_format(paragraph, style_name)
        
        # 遍歷所有表格中的段落
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        style_name = paragraph.style.name if paragraph.style else None
                        if style_name in heading_styles_to_process:
                            apply_custom_format(paragraph, style_name)
        
        # 儲存修改
        doc.save(docx_file)
        print("標題樣式處理完成")
        return True
        
    except Exception as e:
        print(f"\n警告：處理標題樣式時發生錯誤: {e}")
        traceback.print_exc()
        return False

def convert_to_docx(markdown_content, output_file):
    """使用 pypandoc 將 Markdown 轉換為 DOCX"""
    try:
        # 檢查 pandoc 是否已安裝
        try:
            pypandoc.get_pandoc_version()
        except OSError:
            print("未檢測到 Pandoc，正在下載並安裝...")
            pypandoc.download_pandoc()
        
        print(f"\n開始轉換為 DOCX 格式...")
        
        # 先將內容寫入臨時 Markdown 檔案，避免編碼問題
        temp_md = "temp_merged.md"
        with open(temp_md, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        # 設定轉換選項
        extra_args = [
            '--standalone',  # 產生獨立的文件
        ]
        
        # 執行轉換（從檔案轉換）
        pypandoc.convert_file(
            temp_md,
            'docx',
            outputfile=output_file,
            extra_args=extra_args
        )
        
        # 刪除臨時檔案
        if os.path.exists(temp_md):
            os.remove(temp_md)
        
        print(f"\n成功生成 DOCX 檔案：{output_file}")
        print(f"檔案大小：{os.path.getsize(output_file) / 1024:.2f} KB")
        
        # 設定行距為 1.5
        set_line_spacing(output_file, 1.5)
        
        # 處理標題樣式：註解掉此行以保留所有級別的標題作為 PDF 書籤
        # process_docx_headings(output_file)
        
        return True
        
    except Exception as e:
        print(f"\n錯誤：轉換過程發生錯誤: {e}")
        return False

def convert_docx_to_pdf(docx_file):
    """將 DOCX 檔案轉換為 PDF，並確保產生書籤(目錄)"""
    try:
        import win32com.client
        import pythoncom
        
        abs_docx = os.path.abspath(docx_file)
        pdf_file = abs_docx.replace('.docx', '.pdf')
        print(f"\n開始轉換為 PDF 格式：{os.path.basename(pdf_file)}")
        
        # 初始化 COM
        pythoncom.CoInitialize()
        
        # 啟動 Word
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0
        
        # 開啟文檔
        doc = word.Documents.Open(abs_docx, ReadOnly=True)
        
        # 參數定義 (來自 Word VBA)
        # WdExportFormat 17 = wdExportFormatPDF
        # WdExportCreateBookmarks 1 = wdExportCreateHeadingBookmarks (用標題建立書籤)
        
        doc.ExportAsFixedFormat(
            OutputFileName=pdf_file,
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,  # wdExportOptimizeForPrint
            CreateBookmarks=1  # wdExportCreateHeadingBookmarks
        )
        
        doc.Close(SaveChanges=False)
        word.Quit()
        pythoncom.CoUninitialize()
        
        print(f"成功生成 PDF 檔案：{os.path.basename(pdf_file)}")
        print(f"檔案大小：{os.path.getsize(pdf_file) / 1024:.2f} KB")
        
        return True
        
    except Exception as e:
        print(f"\n警告：轉換 PDF 時發生錯誤: {e}")
        traceback.print_exc()
        
        # 如果發生例外，嘗試釋放資源
        try:
            doc.Close(SaveChanges=False)
        except:
            pass
        try:
            word.Quit()
        except:
            pass
        try:
            pythoncom.CoUninitialize()
        except:
            pass
            
        return False

def convert_string_simplified_to_traditional(text):
    """
    使用 Word COM 將字串由簡體轉為繁體。
    需在 Windows 上執行且已安裝 Microsoft Word。
    回傳轉換後的字串；若失敗則回傳原字串。
    增強版：包含重試機制、延遲和完善的資源清理。
    """
    if sys.platform != 'win32':
        return text
    
    import time
    
    # 重試最多 3 次
    for attempt in range(3):
        app = None
        doc = None
        try:
            import pythoncom
            import win32com.client
            
            # 初始化 COM
            pythoncom.CoInitialize()
            
            # 創建 Word 應用程序
            app = win32com.client.Dispatch('Word.Application')
            app.Visible = False
            app.DisplayAlerts = 0  # 禁用所有警告彈窗
            
            # 創建新文檔
            doc = app.Documents.Add()
            doc.Content.Text = text
            
            # 執行簡轉繁
            doc.Content.TCSCConverter(WD_TCSC_CONVERTER_DIRECTION_SCTC, True, True)
            result = doc.Content.Text.strip()
            
            # 關閉文檔（不保存）
            doc.Close(SaveChanges=False)
            doc = None
            
            # 關閉 Word
            app.Quit()
            app = None
            
            # 清理 COM
            pythoncom.CoUninitialize()
            
            return result
            
        except Exception as e:
            print(f"警告：字串簡轉繁失敗 (嘗試 {attempt + 1}/3): {e}")
            
            # 清理資源
            try:
                if doc is not None:
                    doc.Close(SaveChanges=False)
            except:
                pass
            
            try:
                if app is not None:
                    app.Quit()
            except:
                pass
            
            try:
                pythoncom.CoUninitialize()
            except:
                pass
            
            # 如果不是最後一次嘗試，等待 1.5 秒後重試
            if attempt < 2:
                time.sleep(1.5)
            else:
                # 最後一次嘗試失敗，返回原字串
                return text
    
    return text

def convert_docx_simplified_to_traditional(docx_path):
    """
    使用 Word COM 將 DOCX 整份內容由簡體轉為繁體並存檔。
    需在 Windows 上執行且已安裝 Microsoft Word。
    回傳 True 成功，False 失敗。
    增強版：包含重試機制、延遲和完善的資源清理。
    """
    if sys.platform != 'win32':
        print("警告：簡轉繁需在 Windows 上使用 Word COM，已略過")
        return False
    
    import time
    
    abs_path = os.path.abspath(docx_path)
    if not os.path.isfile(abs_path):
        print(f"警告：找不到檔案 {abs_path}")
        return False
    
    # 重試最多 3 次
    for attempt in range(3):
        app = None
        doc = None
        try:
            import pythoncom
            import win32com.client
            
            # 初始化 COM
            pythoncom.CoInitialize()
            
            # 創建 Word 應用程序
            app = win32com.client.Dispatch('Word.Application')
            app.Visible = False
            app.DisplayAlerts = 0  # 禁用所有警告彈窗
            
            # 開啟文檔
            doc = app.Documents.Open(FileName=abs_path, ConfirmConversions=False, ReadOnly=False)
            doc.Activate()
            
            # 選取整份文件後再轉換
            app.Selection.WholeStory()
            app.Selection.Range.TCSCConverter(WD_TCSC_CONVERTER_DIRECTION_SCTC, True, True)
            
            # 保存並關閉
            doc.Save()
            doc.Close(SaveChanges=True)
            doc = None
            
            # 關閉 Word
            app.Quit()
            app = None
            
            # 清理 COM
            pythoncom.CoUninitialize()
            
            print("已使用 Word 將整份文件轉為繁體並存檔")
            return True
            
        except Exception as e:
            print(f"\n警告：Word 簡轉繁時發生錯誤 (嘗試 {attempt + 1}/3): {e}")
            if attempt < 2:
                print(f"將在 1.5 秒後重試...")
            
            # 清理資源
            try:
                if doc is not None:
                    doc.Close(SaveChanges=False)
            except:
                pass
            
            try:
                if app is not None:
                    app.Quit()
            except:
                pass
            
            try:
                pythoncom.CoUninitialize()
            except:
                pass
            
            # 如果不是最後一次嘗試，等待 1.5 秒後重試
            if attempt < 2:
                time.sleep(1.5)
            else:
                # 最後一次嘗試失敗
                traceback.print_exc()
                return False
    
    return False

def main():
    """主程式"""
    print("=" * 60)
    print("Markdown 轉 DOCX/PDF 合併轉換工具")
    print("=" * 60)
    
    # 1. 收集 Markdown 檔案並按書名分組
    print("\n步驟 1: 收集 Markdown 檔案...")
    books = collect_markdown_files_by_book()
    
    if not books:
        print("錯誤：未找到任何符合 書名_數字.md 格式的 Markdown 檔案")
        return 1
    
    print(f"\n找到 {len(books)} 本書，共 {sum(len(files) for files in books.values())} 個檔案：")
    for book_name, files in books.items():
        print(f"\n【{book_name}】：{len(files)} 篇")
        for i, file in enumerate(files, 1):
            article_num = get_article_number(file)
            print(f"  {i}. {file} (第 {article_num} 篇)")
    
    # 2. 處理每本書
    print("\n" + "=" * 60)
    success_count = 0
    fail_count = 0
    
    for book_idx, (book_name, files) in enumerate(books.items(), 1):
        print(f"\n[{book_idx}/{len(books)}] 處理：{book_name}")
        print("-" * 60)
        
        # 合併檔案內容
        print(f"步驟 2: 合併 {len(files)} 個 Markdown 檔案...")
        merged_content = merge_markdown_files(files)
        
        if not merged_content:
            print(f"警告：無法讀取 {book_name} 的內容")
            fail_count += 1
            continue
        
        print(f"合併完成，總字數：{len(merged_content)} 字元")
        
        # 轉換為 DOCX
        print(f"步驟 3: 轉換為 DOCX 格式...")
        output_file = f"{book_name.replace('_', ' ')}.docx"
        
        if convert_to_docx(merged_content, output_file):
            # 使用 Word COM 將 DOCX 簡體轉繁體並存檔
            print(f"步驟 4: 使用 Word 將 DOCX 簡體轉繁體...")
            convert_docx_simplified_to_traditional(output_file)
            
            # 將書名轉為繁體並重新命名檔案
            print(f"步驟 5: 將檔案名稱轉為繁體...")
            traditional_book_name = convert_string_simplified_to_traditional(book_name)
            if traditional_book_name != book_name:
                new_output_file = f"{traditional_book_name.replace('_', ' ')}.docx"
                try:
                    if os.path.exists(new_output_file):
                        os.remove(new_output_file)
                    os.rename(output_file, new_output_file)
                    output_file = new_output_file
                    print(f"已將檔案重新命名為：{output_file}")
                except Exception as e:
                    print(f"警告：重新命名失敗: {e}")
            
            # PDF 由已轉繁的 .docx 產生
            print(f"步驟 6: 轉換為 PDF 格式...")
            convert_docx_to_pdf(output_file)  # PDF 轉換失敗不影響整體成功狀態
            success_count += 1
        else:
            fail_count += 1
    
    # 總結
    print("\n" + "=" * 60)
    print("轉換完成！")
    print(f"成功：{success_count} 本，失敗：{fail_count} 本")
    print("\n已生成 DOCX 和 PDF 檔案")
    print("=" * 60)
    
    return 0 if fail_count == 0 else 1

if __name__ == '__main__':
    sys.exit(main())
